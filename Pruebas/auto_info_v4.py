from docxtpl import DocxTemplate, RichText
from docx import Document
from datetime import datetime
import calendar
import pandas as po
import os
import re
from docx.oxml import OxmlElement
from docx.shared import Pt
from docx.oxml.ns import qn
import win32com.client as win32
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from pathlib import Path


# Funcion para Tabla Prueba
def set_borders(table):

    tbl = table._element
    tblBorders = OxmlElement('w:tblBorders')

    for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = OxmlElement(f"w:{border_name}")
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '8')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    tbl.tblPr.append(tblBorders)

def aplicar_formatos(celda, negrita = False):

    for parrafo in celda.paragraphs:
        for run in parrafo.runs:
            run.font.name = "Century Gothic"
            run.font.size = Pt(11)
            run.bold = negrita

# Funcion #1 para Probar en Plazos


def fucionar_por_plazos (items, marcador = "*", espacio_despues=True, unir_inicio = False):

    resultado = []
    acumulado = False
    buffer = []
    inicio = []

    for token in items:
        if token == marcador:

            if unir_inicio and inicio:
                resultado.append(" ".join(inicio))
                inicio = []
            elif inicio:
                resultado.extend(inicio)
                inicio = []

            if acumulado:
                prefijo = (marcador + " ") if espacio_despues else marcador
                resultado.append(prefijo + " ".join(buffer))
                buffer = []
            acumulado = True

        else:
            if acumulado:
                buffer.append(token)
            else:
                resultado.append(token)
    if acumulado:
        prefijo = (marcador + " ") if espacio_despues else marcador
        resultado.append(prefijo + " ".join(buffer))
    elif inicio:
        if unir_inicio:
            resultado.apend("".join(inicio))
        else:
            resultado.extend(inicio)
    
    return resultado

def separar_periodo (periodo: str):
    a, b = re.split(r"\s*hasta\s*", periodo, flags=re.IGNORECASE)
    return a.strip(), b.strip()

def plazos_unidos(lista_1,lista_2):
    titulos1, periodo1 = lista_1[:-1], lista_1[-1]
    titulos2, periodo2 = lista_2[:-1], lista_2[-1]

    ini1, fin1 = separar_periodo(periodo1)
    ini2, fin2 = separar_periodo(periodo2)

    orden = titulos2 + [t for t in titulos1 if t not in titulos2]

    plazos_por_titulo = {}

    for t in orden:
        if t in titulos1 and t in titulos2:
            plazos_por_titulo[t] = f"{ini1} hasta {fin2}"
        elif t in titulos1:
            plazos_por_titulo[t] = periodo1
        else:
            plazos_por_titulo[t] = periodo2
    return [plazos_por_titulo[t] for t in orden]

# Abrir Excel:

try:
    df = po.read_excel("CTDR_Asistente_Desarrollo_Informatico.xlsx", sheet_name="TDR Asistente v02", engine = "openpyxl")
except FileNotFoundError:
    print("No se encontró el archivo CTDR_Asistente_Desarrollo_Informatico.xlsx")
except Exception as e:
    print("Error al abrir el archivo CTDR_Asistente_Desarrollo_Informatico.xlsx: "+e)

# Se determina cuantos productos existen en el TDR

try:
    numero_inicial_prod = df.iloc[33:,0]
    numeros = numero_inicial_prod[numero_inicial_prod.apply(lambda x: isinstance(x, (int)))]
    ultimo_valor = numeros.iloc[-1] if not numeros.empty else None
except Exception:
    ultimo_valor=0

# Con el último valor es posible recorrer exactamente 
# las actividades y productos para almacenarlos en listas


####Lista de Actividades y productos
actividades_lista, productos_lista = [], []

# Posición Inicial Actividades
columna_actividades = 1
fila_actividades = 33 # Recordemos que usamos ultimo_valor para saber cuantas posiciones recorrer

# Posicion Inicial de los Productos
columna_productos = 6
fila_productos = 33 # Recordemos que usamos ultimo_

# Extraen valores de actividad y productos
if ultimo_valor>0:
    try:
        # Extraen valores de actividad
        serie_actividades = df.iloc[fila_actividades:fila_actividades + ultimo_valor, columna_actividades]
        actividades_lista = serie_actividades.tolist()
    except Exception:
        print("No se pudo extraer los datos de Actividades")
    try:
        # Extraen valores de los Productos
        serie_productos = df.iloc[fila_productos: fila_productos + ultimo_valor, columna_productos]
        productos_lista = serie_productos.tolist()
    except Exception:
        print("No se pudo extraer los datos de Productos")

print(f"Actividades cargadas: {len(actividades_lista)}")
print(f"Productos cargadas: {len(productos_lista)}")


# Una vez creadas las listas, se procede a extraer información importante para el contexto generalizado
# Se extraen los valores del proyecto, puesto y honorarios

try:
    proyecto = df.iloc[9,4].upper()
    puesto = df.iloc[10,4]
    honorario = df.iloc[68,4]
except Exception:
    proyecto, puesto, honorario="No disponible", "No disponible", 0

# Se extraen valores de fechas

fecha_actual = datetime.now()
inicio_periodo = fecha_actual.replace(day=1).strftime("%d-%m-%Y")
ultimo_dia = calendar.monthrange(fecha_actual.year, fecha_actual.month)[1]
fin_periodo = fecha_actual.replace(day=ultimo_dia).strftime("%d-%m-%Y")
meses = {
    1: "enero", 2:"febrero", 3:"marzo", 4:"abril", 5:"mayo", 6:"junio", 7:"julio", 8:"agosto", 9:"septiembre", 10:"octubre", 11:"noviembre", 12:"diciembre"
}
mes_Actual=meses[fecha_actual.month]

##### Plazos extracción ####
metodologia = df.iloc[23,0]
lst_metodologia = metodologia.split()
indices = [i for i, palabra in enumerate(lst_metodologia)
           if palabra ==("Fecha:") ]  #comparar con más posibilidades
plazos = []
for i in indices:
    rango = lst_metodologia[i+2:i+10]
    plazos.append(" ".join(rango))
if len(indices)>2:
    segunda = lst_metodologia[indices[1]+2]
    print(" ".join(segunda))
try:
    plazo_1=plazos[0]
    plazo_2=plazos[1]
    print(f"Plazos cargados: {len(plazos)}")
except Exception:
    print("Error al cargar los plazos")

# Seccion de Fechas para el periodo escrito en texto natural

hoy = datetime.now()
inicio = hoy.replace(day=1)
ultimo = calendar.monthrange(hoy.year, hoy.month)[1]
fin = hoy.replace(day = ultimo)

periodo_incluir_info = f"{inicio.day:02d} al {fin.day:02d} de {meses[hoy.month]} de {hoy.year}"

# Se tiene que idear una forma dinámica de llenar el campo de producto y actividades dependiendo de cuantas actividades o productos existan
# Pero por lo pronto se llenan los otros campos con lo que ya tenemos


# Se crea el contexto para utilizarlo en la herramienta Docx
contexto_plantilla = {
    #"numero": ultimo_valor,
    #"producto": productos_lista, #Especial ciudado con este campo
    "proyecto": proyecto,
    "mes": mes_Actual.upper(),
    "funcionario": "BRYAN BENJAMÍN SARABINO CUICHÁN",
    "cedula": "172231964-5",
    "puesto": puesto,
    "honorario":honorario,
    "fecha_inicio":inicio_periodo,
    "fecha_fin":fin_periodo,
    "periodo":periodo_incluir_info,
    #"actividad":actividades_lista, #Especial cuidado con este campo
    "conclusion_1":"Escribir conclusión #1",
    "conclusion_2":"Escribir conclusión #2",
    "conclusion_3":"Escribir conclusión #3",
    "recomendacion_1":"Escribir recomendación #1",
    "recomendacion_2":"Escribir recomendación #2",
    "recomendacion_3":"Escribir recomendación #3"

}

# Una vez hemos creado el contexto general para la 
# plantilla de productos, se tiene que proceder 
# a llenar los campos


##############################################################################################

# Sección de Pruebas

# Se abre el archivo de excel con el path para las plantillas
try:
    df_pl = po.read_excel("PLANTILLAS_INFORMES.xlsx", sheet_name="Hoja1", engine = "openpyxl")
except FileNotFoundError:
    print("El arichivo no PLANTILLAS_INFORMES.xlsx no se encontró.")
except Exception as e:
    print("Error al abrir el archivo PLANTILLAS_INFORMES.xlsx.")
# Ahora se realiza una prueba con un menu simple (FUNCIONAL PARA LOS PRODUCTOS)

bandera = "s"
while bandera == "s":
    print(" -------------------- Menú principal -------------------")
    print("1. Informe de Productos")
    print("2. Informe de Actividades y Productos Entregados")
    print("3. Informe de Aceptación de los Productos Recibidos a Satisfacción")

    seleccion_menu = input("Indique el informe (1 - 3): ")
    if (seleccion_menu == "1"):
        if ultimo_valor==0:
            print("No hay productos para generar")
        else:
            print("\nElija cual informe realizar")
            for i in range(ultimo_valor):
                print(f"{i + 1}. Producto {i + 1}")
            print(f"{ultimo_valor + 1}. Generar todos los Informes")

            usuario_seleccion = input(f"Digite su elección (1 - {ultimo_valor+1}): ")
            if not usuario_seleccion.isdigit():
                print("Debe ingresar un número")
            else:
                usuario_seleccion = int(usuario_seleccion)
                if 1<= usuario_seleccion <= ultimo_valor:
                    try:
                        path_informe = df_pl.iloc[0,2]
                        plantilla_prod = DocxTemplate(path_informe)
                        contexto_actividad_prod = {
                            **contexto_plantilla,
                            "producto":productos_lista[int(usuario_seleccion) - 1],
                            "actividad":actividades_lista[int(usuario_seleccion) - 1],
                            "numero": usuario_seleccion
                        }
                        plantilla_prod.render(contexto_actividad_prod)
                        plantilla_prod.save(f"Informe_de_Producto_{usuario_seleccion}.docx")
                    except Exception as e:
                        print(f"Error al generar el informe: {e}")

                    ruta_absoluta_pro = os.path.abspath(f"Informe_de_Producto_{usuario_seleccion}.docx")
                    directorio = os.path.dirname(ruta_absoluta_pro)
                    print(f"Revise su directorio: {directorio}")
                    print()
                elif usuario_seleccion == ultimo_valor+1:
                    for i in range(1, ultimo_valor+1):
                        try:
                            path_informe = df_pl.iloc[0,2]
                            plantilla_prod = DocxTemplate(path_informe)
                            contexto_actividad_prod = {
                                **contexto_plantilla,
                                "producto":productos_lista[i - 1],
                                "actividad":actividades_lista[i - 1],
                                "numero": i
                            }
                            plantilla_prod.render(contexto_actividad_prod)
                            plantilla_prod.save(f"Informe_de_Producto_{i}.docx")
                        except Exception as e:
                            print(f"Error en el informe {i}: {e}")
                    ruta_absoluta_produ = os.path.abspath(f"Informe_de_Producto_{i}.docx")
                    directorio = os.path.dirname(ruta_absoluta_produ)
                    print(f"Revise su directorio: {directorio}")
                    print()

                else:
                    print("Error fuera de rango.")

    elif (seleccion_menu == "2"):
        print()
        # CAMBIOS NECESARIOS: 
        # HACER QUE EL USUARIO PUEDA INGRESAR UNO O MÁS PRODUCTOS / ACTIVIDADES AL INFORME
        # OPCION DE MEJORA: PERMITIR QUE SE INGRESE MÁS DE UN NÚMERO SEPARADO POR ESPACIOS
        #                   Y LUEGO HAGA UN SPLIT Y GUARDE ESO POR COMAS EN UNA LISTA

        print()
        print("Indique cuantos productos y/o actividades tiene el informe:")
        for i in range(ultimo_valor):
            print(f"{i + 1}. Producto / Actividad {i + 1}")
            print()
        print("Digite la letra 'T' para generar un informe con todos los Productos / Actividades realizadas.")
        print()
        print("Indique de que Producto / Actividad realizar el reporte.")
        print("Recuerde que puede digitar más de una opción, separelo por espacios.")
        usuario_seleccion = input(f"Ingrese sus selecciones (1 - {ultimo_valor} / T): ")

        if (usuario_seleccion != 'T'):
            lista_seleccion = usuario_seleccion.split()

            # Control de Usuarios
            # Se controla que las opciones seleccionadas no superen al último valor
            lista_depurada = []
            for i in lista_seleccion:
                if ((int(i) >= ultimo_valor + 1)):
                    print("Lo sentimos, a ingresado un valor invalido.")
                    print(f"No existe actualmente un Producto / Actividad {i}")
                    print("Se removerá a esta opción de su selección.")
                else:
                    lista_depurada.append(int(i))

            productos_seleccionados = []
            actividades_seleccionadas = []

            for i in lista_depurada:
                productos_seleccionados.append(productos_lista[i - 1])
                actividades_seleccionadas.append(actividades_lista[i -1])
            print()
            print("A continuación se generará su informe")
            df_pl = po.read_excel("PLANTILLAS_INFORMES.xlsx", sheet_name="Hoja1", engine = "openpyxl")
            path_informe = df_pl.iloc[1,2]

            plantilla_prod_act = DocxTemplate(path_informe)

            contexto_actividad_prod = {
                **contexto_plantilla,
                "productos_lista":productos_seleccionados,
                "actividades_lista":actividades_seleccionadas,
            }

            plantilla_prod_act.render(contexto_actividad_prod)
            plantilla_prod_act.save("Informe_de_Actividad_Producto_Realizados_Total.docx")

            print("Revise su directorio")
            print()
        else:
            print()
            df_pl = po.read_excel("PLANTILLAS_INFORMES.xlsx", sheet_name="Hoja1", engine = "openpyxl")
            path_informe = df_pl.iloc[1,2]

            plantilla_prod_act = DocxTemplate(path_informe)

            contexto_actividad_prod = {
                **contexto_plantilla,
                "productos_lista":productos_lista,
                "actividades_lista":actividades_lista,
            }
            plantilla_prod_act.render(contexto_actividad_prod)
            plantilla_prod_act.save("Informe_de_Actividad_Producto_Realizados_Total.docx")

            print("Revise su directorio")
            print()

    elif (seleccion_menu == "3"):

        PLACEHOLDER = "[[TABLA_PRODUCTOS]]"

        plazos_de_informe = []

        print()
        print("Indique cuantos productos tiene el informe:")
        for i in range(ultimo_valor):
            print(f"{i + 1}. Producto {i + 1}")
            print()
        print("Digite la letra 'T' para generar un informe con todos los Productos / Actividades realizadas.")
        print()
        print("Indique de que Producto realizará el reporte.")
        print("Recuerde que puede digitar más de una opción, separelo por espacios.")
        usuario_seleccion = input(f"Ingrese sus selecciones (1 - {ultimo_valor} / T): ")

        if (usuario_seleccion != 'T'):
            lista_seleccion = usuario_seleccion.split()

            # Control de Usuarios
            # Se controla que las opciones seleccionadas no superen al último valor
            lista_depurada = []
            for i in lista_seleccion:
                if ((int(i) >= ultimo_valor + 1)):
                    print("Lo sentimos, a ingresado un valor invalido.")
                    print(f"No existe actualmente un Producto {i}")
                    print("Se removerá a esta opción de su selección.")
                else:
                    lista_depurada.append(int(i))

            productos_seleccionados = []

            for i in lista_depurada:
                productos_seleccionados.append(productos_lista[i - 1])
            

            #PLAZOS PRODUCTOS

            ##### Plazos extracción ####
            listas_para_unir = []
            metodologia = df.iloc[23,0]
            lst_metodologia = metodologia.split()
            indice_1 = [i for i, palabra in enumerate(lst_metodologia)
                    if palabra ==("Fecha:") ]  #comparar con más posibilidades
                                               # utilizar solo el primer índice
            
            indice_2 = [i for i, palabra in enumerate(lst_metodologia)
                        if palabra == ("Honorarios:")] # utilizar solo el primer índice

            lista_plazos = lst_metodologia[indice_1[0]:indice_2[0]]
            
            plazos_separados = []
            actual = []
            for elemento in lista_plazos:
                if elemento == "Fecha:":

                    if actual:
                        plazos_separados.append(actual)
                    actual = [elemento]
                else:
                    actual.append(elemento)
            if actual:
                plazos_separados.append(actual)
            for plazos in plazos_separados:
                plazos_semi_procesados = fucionar_por_plazos(plazos, marcador = "*",espacio_despues=True ,unir_inicio=True)
                
                informes_plazo = plazos_semi_procesados[10:].copy()
                plazos_texto = ""

                for i in range(2, 10):
                    plazos_texto += plazos_semi_procesados[i] + " "
                
                plazos_texto = plazos_texto.strip()
                informes_plazo.append(plazos_texto)
                listas_para_unir.append(informes_plazo)

            if len(listas_para_unir) >= 2:
                lista1, lista2 = listas_para_unir[0], listas_para_unir[1]
                plazos_final = plazos_unidos(lista1, lista2)
            else:
                plazos_final = listas_para_unir[0] if listas_para_unir else []

            print(productos_lista)
            print(plazos_final)

            doc = Document()

            tabla = doc.add_table(rows=3, cols=3)

            
            headers = ["Nombre", "Productos que deben ser entregados por el personal contratado", "Plazo de entrega de los productos"]

            for i, texto in enumerate(headers):
                celda = tabla.rows[0].cells[i]
                celda.text = texto
                aplicar_formatos(celda, negrita = True)

            top_cell = tabla.cell(1,0)
            bottom_cell = tabla.cell(2,0)
            top_cell.merge(bottom_cell)
            top_cell.text = "BRYAN BENJAMIN SARABINO CUICHÁN"
            top_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            aplicar_formatos(top_cell)

            tabla.cell(1,1).text = productos_lista[0]
            tabla.cell(1,2).text = plazos_final[0]
            tabla.cell(2,1).text = productos_lista[1]
            tabla.cell(2,2).text = plazos_final[1]

            for i in range(1,3):
                for j in range(1,3):
                    aplicar_formatos(tabla.cell(i,j))
              
            set_borders(tabla)
            doc.save("tabla_base.docx")

            df_pl = po.read_excel("PLANTILLAS_INFORMES.xlsx", sheet_name="Hoja1", engine = "openpyxl")
            path_informe = df_pl.iloc[2,2]
            contexto_actividad_prod = {
                **contexto_plantilla,
                "productos_lista":productos_lista,
                "actividades_lista":actividades_lista,
            }
            plantilla_prod_acpt = DocxTemplate(path_informe)
            subdoc = plantilla_prod_acpt.new_subdoc("tabla_base.docx")
            contexto = {"tabla": subdoc}
            plantilla_prod_acpt.render(contexto)
            plantilla_prod_acpt.save("PRUEBA_PROD_ACT_V2.docx")

            print("Revise su directorio")
            print()


    bandera = input("¿Desea generar otro informe? (s/n): ")

    print("¡Muchas gracias por usar el programa!")
