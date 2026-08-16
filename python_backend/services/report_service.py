from pathlib import Path
from docx import Document

def generar_word_cedula(cedula:str, info: dict | None = None) -> Path:
    out_dir = Path(__file__).resolve().parent.parent / "out"
    out_dir.mkdir(exist_ok=True)

    filename = f"Reporte_{cedula}.docx"
    out_path = out_dir / filename

    doc = Document()
    doc.add_heading("Reporte de Funcionario", level=1)
    doc.add_paragraph(f"Cedula: {cedula}")

    if info:
        doc.add_paragraph(f"Nombres: {info.get('nombres','')}")
        doc.add_paragraph(f"Apellidos: {info.get('apellidos','')}")
        doc.add_paragraph(f"Dirección: {info.get('direccion','')}")
        doc.add_paragraph(f"Cargo: {info.get('cargo','')}")
    doc.add_paragraph("")
    doc.save(out_path)
    return out_path